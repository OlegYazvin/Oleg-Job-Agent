from datetime import datetime
from pathlib import Path

from docx import Document

from job_agent.models import (
    FirstOrderMessage,
    JobOutreachBundle,
    JobPosting,
    ManualReviewLink,
    NearMissJob,
    SecondOrderIntroMessage,
)
from job_agent.reports import (
    build_live_outreach_payload,
    build_message_document,
    build_near_miss_document,
    build_near_miss_payload,
    build_summary_document,
)


def build_bundle() -> JobOutreachBundle:
    job = JobPosting(
        company_name="Acme AI",
        role_title="Senior Product Manager, AI",
        direct_job_url="https://boards.greenhouse.io/acme/jobs/123",
        ats_platform="Greenhouse",
        location_text="Remote",
        is_fully_remote=True,
        posted_date_text="2 days ago",
        posted_date_iso=None,
        base_salary_min_usd=210000,
        base_salary_max_usd=240000,
        salary_text="$210,000 - $240,000",
        evidence_notes="Direct ATS, remote, recent, salary listed.",
    )
    return JobOutreachBundle(
        job=job,
        first_order_messages=[
            FirstOrderMessage(
                contact_name="Jane Doe",
                contact_profile_url="https://www.linkedin.com/in/jane-doe/",
                subject_context="Past collaborator",
                message_body="Would you be open to connecting me with the hiring team?",
            )
        ],
        second_order_messages=[
            SecondOrderIntroMessage(
                first_order_contact_name="John Smith",
                first_order_contact_profile_url="https://www.linkedin.com/in/john-smith/",
                second_order_contact_names=["Priya Patel", "Alex Chen"],
                second_order_contact_profile_urls=[
                    "https://www.linkedin.com/in/priya-patel/",
                    "https://www.linkedin.com/in/alex-chen/",
                ],
                message_body="Would you be open to introducing me to Priya Patel?",
            )
        ],
    )


def test_reports_are_written(tmp_path: Path) -> None:
    bundle = build_bundle()
    generated_at = datetime(2026, 3, 25, 14, 44, 53)
    message_doc = build_message_document([bundle], tmp_path, generated_at=generated_at)
    summary_doc = build_summary_document([bundle], tmp_path, generated_at=generated_at)
    assert message_doc.exists()
    assert summary_doc.exists()
    assert message_doc.name == "linkedin_outreach_messages-20260325-144453.docx"
    assert summary_doc.name == "job_summary-20260325-144453.docx"
    assert (tmp_path / "linkedin_outreach_messages.docx").exists()
    assert (tmp_path / "job_summary.docx").exists()


def test_message_report_includes_manual_review_links(tmp_path: Path) -> None:
    bundle = build_bundle().model_copy(
        update={
            "first_order_messages": [],
            "second_order_messages": [],
            "manual_review_links": [
                ManualReviewLink(
                    label="1st-degree people at company",
                    url="https://www.linkedin.com/search/results/people/?keywords=Acme%20AI&network=%5B%22F%22%5D",
                )
            ],
            "manual_review_notes": ["Use Current Company filter in LinkedIn."],
        }
    )
    message_doc = build_message_document([bundle], tmp_path)
    doc = Document(message_doc)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Manual LinkedIn review mode is enabled" in text
    assert "1st-degree people at company" in text


def test_live_outreach_payload_includes_live_message_metadata() -> None:
    bundle = build_bundle()
    payload = build_live_outreach_payload([bundle], run_id="run-123", generated_at=datetime(2026, 3, 25, 14, 44, 53))
    assert payload["run_id"] == "run-123"
    assert payload["item_count"] == 2
    assert payload["items"][0]["recipient_name"] == "Jane Doe"
    assert payload["items"][0]["live"] is True
    assert payload["items"][1]["target_names"] == ["Priya Patel", "Alex Chen"]


def test_summary_report_counts_second_order_targets_not_messages(tmp_path: Path) -> None:
    bundle = build_bundle()
    summary_doc = build_summary_document([bundle], tmp_path)
    doc = Document(summary_doc)
    table = doc.tables[0]
    assert table.rows[0].cells[7].text == "2nd-Degree Contacts Messaged"
    assert table.rows[1].cells[7].text == "2"


def test_summary_report_renders_still_open_again_section_for_reacquired_jobs(tmp_path: Path) -> None:
    bundle = build_bundle()
    reacquired_job = JobPosting(
        company_name="Bravo AI",
        role_title="Principal Product Manager, AI",
        direct_job_url="https://jobs.lever.co/bravo/123",
        resolved_job_url="https://jobs.lever.co/bravo/123",
        ats_platform="Lever",
        location_text="Remote",
        is_fully_remote=True,
        posted_date_text="2026-03-28",
        posted_date_iso="2026-03-28",
        base_salary_min_usd=220000,
        base_salary_max_usd=260000,
        salary_text="$220,000 - $260,000",
        evidence_notes="Still live and fully qualified.",
        is_reacquired=True,
        first_reported_at="2026-03-20T10:00:00+00:00",
        last_reported_at="2026-03-28T10:00:00+00:00",
        report_count=3,
    )

    summary_doc = build_summary_document([bundle], tmp_path, reacquired_jobs=[reacquired_job])
    doc = Document(summary_doc)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "Still Open Again" in text
    assert "Bravo AI | Principal Product Manager, AI" in text
    assert "First reported: 2026-03-20T10:00:00+00:00" in text


def test_near_miss_report_and_payload_are_written(tmp_path: Path) -> None:
    near_miss = NearMissJob(
        company_name="Acme AI",
        role_title="Staff Product Manager, AI",
        reason_code="missing_salary",
        detail="No salary range was published on the direct page.",
        why_close="The role looked strong on a real ATS page, but compensation was missing.",
        source_url="https://boards.greenhouse.io/acme/jobs/123",
        direct_job_url="https://boards.greenhouse.io/acme/jobs/123",
        source_type="direct_ats",
        posted_date_text="2026-03-25",
        is_remote=True,
        supporting_evidence=["Remote ATS role."],
    )
    generated_at = datetime(2026, 3, 25, 14, 44, 53)

    near_miss_doc = build_near_miss_document([near_miss], tmp_path, generated_at=generated_at)
    payload = build_near_miss_payload([near_miss], run_id="run-123", generated_at=generated_at)

    assert near_miss_doc.exists()
    assert near_miss_doc.name == "job_near_misses-20260325-144453.docx"
    assert payload["near_miss_count"] == 1
    assert payload["items"][0]["reason_code"] == "missing_salary"
