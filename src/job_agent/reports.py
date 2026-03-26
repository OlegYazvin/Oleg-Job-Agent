from __future__ import annotations

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document

from .models import JobOutreachBundle, RunManifest


def _add_paragraph_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        document.add_paragraph(line)


def _count_second_order_targets_with_messages(bundle: JobOutreachBundle) -> int:
    profile_urls: set[str] = set()
    for message in bundle.second_order_messages:
        profile_urls.update(message.second_order_contact_profile_urls)
    return len(profile_urls)


def _timestamp_slug(generated_at: datetime) -> str:
    return generated_at.strftime("%Y%m%d-%H%M%S")


def _write_latest_alias(source_path: Path, latest_path: Path) -> None:
    if source_path == latest_path:
        return
    shutil.copy2(source_path, latest_path)


def build_message_document(
    bundles: list[JobOutreachBundle],
    output_dir: Path,
    *,
    generated_at: datetime | None = None,
) -> Path:
    generated_at = generated_at or datetime.now()
    document = Document()
    document.add_heading("LinkedIn Outreach Drafts", level=0)
    document.add_paragraph(f"Generated at: {generated_at.isoformat(timespec='seconds')}")

    if not bundles:
        document.add_paragraph("No jobs matched the configured criteria in this run.")

    for bundle in bundles:
        job = bundle.job
        document.add_heading(f"{job.company_name} | {job.role_title}", level=1)
        if not bundle.first_order_messages and not bundle.second_order_messages:
            if bundle.manual_review_links:
                document.add_paragraph(
                    "Manual LinkedIn review mode is enabled for this job. Use the links below to find 1st/2nd-degree contacts."
                )
                for link in bundle.manual_review_links:
                    document.add_paragraph(f"{link.label}: {link.url}")
                for note in bundle.manual_review_notes:
                    document.add_paragraph(f"Note: {note}")
            else:
                document.add_paragraph("No LinkedIn messages were drafted for this job.")
            continue

        if bundle.first_order_messages:
            document.add_paragraph("Direct company contacts")
        for message in bundle.first_order_messages:
            document.add_heading(f"To: {message.contact_name}", level=2)
            document.add_paragraph(f"Recipient profile: {message.contact_profile_url}")
            document.add_paragraph(message.message_body)

        if bundle.second_order_messages:
            document.add_paragraph("Warm intro paths")
        for message in bundle.second_order_messages:
            document.add_heading(f"To: {message.first_order_contact_name}", level=2)
            if message.first_order_contact_profile_url:
                document.add_paragraph(f"Recipient profile: {message.first_order_contact_profile_url}")
            if message.second_order_contact_names:
                document.add_paragraph(f"Targets: {', '.join(message.second_order_contact_names)}")
            document.add_paragraph(message.message_body)

    path = output_dir / f"linkedin_outreach_messages-{_timestamp_slug(generated_at)}.docx"
    document.save(path)
    _write_latest_alias(path, output_dir / "linkedin_outreach_messages.docx")
    return path


def build_summary_document(
    bundles: list[JobOutreachBundle],
    output_dir: Path,
    *,
    generated_at: datetime | None = None,
) -> Path:
    generated_at = generated_at or datetime.now()
    document = Document()
    document.add_heading("Job Summary", level=0)
    document.add_paragraph(f"Generated at: {generated_at.isoformat(timespec='seconds')}")

    if not bundles:
        document.add_paragraph("No qualifying jobs were found in this run.")

    table = document.add_table(rows=1, cols=8)
    headers = table.rows[0].cells
    headers[0].text = "Company"
    headers[1].text = "Role"
    headers[2].text = "Job Link"
    headers[3].text = "Posted"
    headers[4].text = "Salary"
    headers[5].text = "1st-Degree Contacts"
    headers[6].text = "2nd-Degree Contacts"
    headers[7].text = "2nd-Degree Contacts Messaged"

    for bundle in bundles:
        row = table.add_row().cells
        row[0].text = bundle.job.company_name
        row[1].text = bundle.job.role_title
        row[2].text = str(bundle.job.direct_job_url)
        row[3].text = bundle.job.posted_date_text
        row[4].text = bundle.job.salary_text or ""
        row[5].text = str(len(bundle.first_order_contacts))
        row[6].text = str(len(bundle.second_order_contacts))
        row[7].text = str(_count_second_order_targets_with_messages(bundle))

    path = output_dir / f"job_summary-{_timestamp_slug(generated_at)}.docx"
    document.save(path)
    _write_latest_alias(path, output_dir / "job_summary.docx")
    return path


def _normalize_name(value: str) -> str:
    return "".join(char.lower() for char in value if char.isalnum())


def _build_history_lookup(bundle: JobOutreachBundle) -> tuple[dict[str, list[str]], dict[str, list[str]]]:
    histories_by_name: dict[str, list[str]] = {}
    histories_by_profile_url: dict[str, list[str]] = {}

    for contact in bundle.first_order_contacts:
        if contact.message_history:
            histories_by_name[_normalize_name(contact.name)] = list(contact.message_history)
            histories_by_profile_url[str(contact.profile_url)] = list(contact.message_history)

    for contact in bundle.second_order_contacts:
        for connector_name, history in contact.connected_first_order_message_histories.items():
            if history:
                histories_by_name.setdefault(_normalize_name(connector_name), list(history))
        for connector_name, profile_url in contact.connected_first_order_profile_urls.items():
            history = contact.connected_first_order_message_histories.get(connector_name)
            if history and profile_url:
                histories_by_profile_url.setdefault(str(profile_url), list(history))

    return histories_by_name, histories_by_profile_url


def build_live_outreach_payload(
    bundles: list[JobOutreachBundle],
    *,
    run_id: str,
    generated_at: datetime | None = None,
) -> dict[str, object]:
    generated_at = generated_at or datetime.now()
    items: list[dict[str, object]] = []
    priority_rank = 1

    for bundle in bundles:
        histories_by_name, histories_by_profile_url = _build_history_lookup(bundle)

        for message in bundle.first_order_messages:
            profile_url = str(message.contact_profile_url)
            history = histories_by_profile_url.get(profile_url) or histories_by_name.get(
                _normalize_name(message.contact_name),
                [],
            )
            items.append(
                {
                    "priority_rank": priority_rank,
                    "company_name": bundle.job.company_name,
                    "role_title": bundle.job.role_title,
                    "job_url": str(bundle.job.direct_job_url),
                    "draft_type": "first_order_info",
                    "recipient_name": message.contact_name,
                    "recipient_profile_url": profile_url,
                    "target_names": [],
                    "target_profile_urls": [],
                    "message_body": message.message_body,
                    "message_history_count": len(history),
                    "live": bool(profile_url),
                }
            )
            priority_rank += 1

        for message in bundle.second_order_messages:
            profile_url = str(message.first_order_contact_profile_url or "")
            history = histories_by_profile_url.get(profile_url) or histories_by_name.get(
                _normalize_name(message.first_order_contact_name),
                [],
            )
            items.append(
                {
                    "priority_rank": priority_rank,
                    "company_name": bundle.job.company_name,
                    "role_title": bundle.job.role_title,
                    "job_url": str(bundle.job.direct_job_url),
                    "draft_type": "second_order_intro",
                    "recipient_name": message.first_order_contact_name,
                    "recipient_profile_url": profile_url or None,
                    "target_names": list(message.second_order_contact_names),
                    "target_profile_urls": list(message.second_order_contact_profile_urls),
                    "message_body": message.message_body,
                    "message_history_count": len(history),
                    "live": bool(profile_url),
                }
            )
            priority_rank += 1

    return {
        "run_id": run_id,
        "generated_at": generated_at.isoformat(timespec="seconds"),
        "item_count": len(items),
        "items": items,
    }


def build_manifest(
    *,
    run_id: str,
    bundles: list[JobOutreachBundle],
    jobs_found_by_search: int,
    message_docx_path: Path,
    summary_docx_path: Path,
    generated_at: datetime | None = None,
) -> RunManifest:
    generated_at = generated_at or datetime.now()
    return RunManifest(
        run_id=run_id,
        generated_at=generated_at,
        message_docx_path=str(message_docx_path),
        summary_docx_path=str(summary_docx_path),
        jobs_found_by_search=jobs_found_by_search,
        jobs_kept_after_validation=len(bundles),
        jobs_with_any_messages=sum(
            1 for bundle in bundles if bundle.first_order_messages or bundle.second_order_messages
        ),
    )
